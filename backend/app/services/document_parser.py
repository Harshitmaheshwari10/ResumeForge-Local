"""Document text extraction from PDF and DOCX."""

import io
import logging
from pathlib import Path

import pdfplumber
from docx import Document
from PyPDF2 import PdfReader

from app.utils.text_utils import normalize_text

logger = logging.getLogger(__name__)


def extract_from_pdf(file_bytes: bytes) -> str:
  """Extract text from PDF using pdfplumber with PyPDF2 fallback."""
  text_parts: list[str] = []

  try:
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
      for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
          text_parts.append(page_text)
  except Exception as exc:
    logger.warning("pdfplumber failed: %s; trying PyPDF2", exc)
    try:
      reader = PdfReader(io.BytesIO(file_bytes))
      for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
          text_parts.append(page_text)
    except Exception as inner:
      raise ValueError(f"Failed to extract PDF text: {inner}") from inner

  if not text_parts:
    raise ValueError("No text could be extracted from the PDF")

  return normalize_text("\n".join(text_parts))


def extract_from_docx(file_bytes: bytes) -> str:
  """Extract text from DOCX including paragraphs and tables."""
  try:
    doc = Document(io.BytesIO(file_bytes))
  except Exception as exc:
    raise ValueError(f"Failed to read DOCX file: {exc}") from exc

  parts: list[str] = []
  for para in doc.paragraphs:
    if para.text.strip():
      parts.append(para.text)

  for table in doc.tables:
    for row in table.rows:
      row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
      if row_text:
        parts.append(row_text)

  if not parts:
    raise ValueError("No text could be extracted from the DOCX")

  return normalize_text("\n".join(parts))


def extract_from_txt(file_bytes: bytes) -> str:
  for encoding in ("utf-8", "latin-1", "cp1252"):
    try:
      return normalize_text(file_bytes.decode(encoding))
    except UnicodeDecodeError:
      continue
  raise ValueError("Could not decode text file")


def extract_text(file_bytes: bytes, filename: str) -> str:
  """Route extraction based on file extension."""
  ext = Path(filename).suffix.lower()
  if ext == ".pdf":
    return extract_from_pdf(file_bytes)
  if ext == ".docx":
    return extract_from_docx(file_bytes)
  if ext in (".txt", ".text"):
    return extract_from_txt(file_bytes)
  raise ValueError(f"Unsupported file type: {ext}")
