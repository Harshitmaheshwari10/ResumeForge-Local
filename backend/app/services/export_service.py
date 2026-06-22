"""Export optimized resume to PDF and DOCX."""

import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.schemas.schemas import ResumeSection
from app.services.resume_rewriter import render_resume_text


def export_to_docx(sections: ResumeSection, output_path: Path) -> Path:
  """Generate ATS-safe DOCX (single column, no tables/icons)."""
  doc = Document()
  sections_margin = doc.sections[0]
  sections_margin.top_margin = Inches(0.6)
  sections_margin.bottom_margin = Inches(0.6)
  sections_margin.left_margin = Inches(0.75)
  sections_margin.right_margin = Inches(0.75)

  contact = sections.contact
  if contact.get("name"):
    name_para = doc.add_paragraph()
    run = name_para.add_run(contact["name"])
    run.bold = True
    run.font.size = Pt(16)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

  contact_line = " | ".join(
    contact[k] for k in ("email", "phone", "linkedin", "website") if contact.get(k)
  )
  if contact_line:
    cp = doc.add_paragraph(contact_line)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cp.runs:
      run.font.size = Pt(10)
      run.font.color.rgb = RGBColor(80, 80, 80)

  doc.add_paragraph()

  def add_section(title: str, content_fn) -> None:
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(30, 30, 30)
    content_fn()
    doc.add_paragraph()

  if sections.summary:
    add_section("PROFESSIONAL SUMMARY", lambda: doc.add_paragraph(sections.summary))

  if sections.skills:
    add_section("SKILLS", lambda: doc.add_paragraph(", ".join(sections.skills)))

  if sections.experience:
    def add_exp() -> None:
      for exp in sections.experience:
        if exp.get("title"):
          p = doc.add_paragraph()
          r = p.add_run(exp["title"] + (f" | {exp['dates']}" if exp.get("dates") else ""))
          r.bold = True
        for bullet in exp.get("bullets", []):
          doc.add_paragraph(bullet, style="List Bullet")

    add_section("PROFESSIONAL EXPERIENCE", add_exp)

  if sections.projects:
    def add_proj() -> None:
      for proj in sections.projects:
        if proj.get("name"):
          p = doc.add_paragraph()
          p.add_run(proj["name"]).bold = True
        if proj.get("description"):
          doc.add_paragraph(proj["description"], style="List Bullet")

    add_section("PROJECTS", add_proj)

  if sections.education:
    def add_edu() -> None:
      for edu in sections.education:
        if edu.get("institution"):
          p = doc.add_paragraph()
          p.add_run(edu["institution"]).bold = True
        if edu.get("degree"):
          doc.add_paragraph(edu["degree"])

    add_section("EDUCATION", add_edu)

  if sections.certifications:
    add_section(
      "CERTIFICATIONS",
      lambda: [doc.add_paragraph(c, style="List Bullet") for c in sections.certifications],
    )

  if sections.achievements:
    add_section(
      "ACHIEVEMENTS",
      lambda: [doc.add_paragraph(a, style="List Bullet") for a in sections.achievements],
    )

  doc.save(str(output_path))
  return output_path


def export_to_pdf(sections: ResumeSection, output_path: Path) -> Path:
  """Generate ATS-safe PDF (single column, standard fonts)."""
  doc = SimpleDocTemplate(
    str(output_path),
    pagesize=letter,
    leftMargin=0.75 * inch,
    rightMargin=0.75 * inch,
    topMargin=0.6 * inch,
    bottomMargin=0.6 * inch,
  )

  styles = getSampleStyleSheet()
  name_style = ParagraphStyle(
    "Name",
    parent=styles["Normal"],
    fontSize=16,
    leading=20,
    alignment=TA_CENTER,
    spaceAfter=4,
    fontName="Helvetica-Bold",
  )
  contact_style = ParagraphStyle(
    "Contact",
    parent=styles["Normal"],
    fontSize=9,
    leading=12,
    alignment=TA_CENTER,
    spaceAfter=12,
    textColor="#505050",
  )
  section_style = ParagraphStyle(
    "Section",
    parent=styles["Normal"],
    fontSize=11,
    leading=14,
    spaceBefore=10,
    spaceAfter=4,
    fontName="Helvetica-Bold",
  )
  body_style = ParagraphStyle(
    "Body",
    parent=styles["Normal"],
    fontSize=10,
    leading=14,
    alignment=TA_LEFT,
    spaceAfter=4,
  )
  bullet_style = ParagraphStyle(
    "Bullet",
    parent=body_style,
    leftIndent=12,
    bulletIndent=0,
    spaceAfter=2,
  )

  story: list = []
  contact = sections.contact

  if contact.get("name"):
    story.append(Paragraph(contact["name"], name_style))

  contact_line = " | ".join(
    contact[k] for k in ("email", "phone", "linkedin", "website") if contact.get(k)
  )
  if contact_line:
    story.append(Paragraph(contact_line, contact_style))

  def add_section(title: str, lines: list[str]) -> None:
    story.append(Paragraph(title, section_style))
    for line in lines:
      story.append(Paragraph(line, bullet_style if line.startswith("•") else body_style))

  if sections.summary:
    add_section("PROFESSIONAL SUMMARY", [sections.summary])

  if sections.skills:
    add_section("SKILLS", [", ".join(sections.skills)])

  if sections.experience:
    exp_lines: list[str] = []
    for exp in sections.experience:
      if exp.get("title"):
        exp_lines.append(f"<b>{exp['title']}</b>" + (f" | {exp['dates']}" if exp.get("dates") else ""))
      for bullet in exp.get("bullets", []):
        exp_lines.append(f"• {bullet}")
    add_section("PROFESSIONAL EXPERIENCE", exp_lines)

  if sections.projects:
    proj_lines: list[str] = []
    for proj in sections.projects:
      if proj.get("name"):
        proj_lines.append(f"<b>{proj['name']}</b>")
      if proj.get("description"):
        proj_lines.append(f"• {proj['description']}")
    add_section("PROJECTS", proj_lines)

  if sections.education:
    edu_lines: list[str] = []
    for edu in sections.education:
      if edu.get("institution"):
        edu_lines.append(f"<b>{edu['institution']}</b>")
      if edu.get("degree"):
        edu_lines.append(edu["degree"])
    add_section("EDUCATION", edu_lines)

  if sections.certifications:
    add_section("CERTIFICATIONS", [f"• {c}" for c in sections.certifications])

  if sections.achievements:
    add_section("ACHIEVEMENTS", [f"• {a}" for a in sections.achievements])

  doc.build(story)
  return output_path


def export_from_text(text: str, output_path: Path, fmt: str) -> Path:
  """Fallback text-based export."""
  from app.services.resume_parser import parse_resume

  parsed = parse_resume(text)
  if fmt == "pdf":
    return export_to_pdf(parsed.sections, output_path)
  return export_to_docx(parsed.sections, output_path)
